from flask import Flask, request, render_template, redirect, send_file
from docx import Document
from docx.shared import Inches
from builtins import str
import os
import docx
from datetime import datetime


app = Flask(__name__)



@app.route('/')
def my_form():
    return render_template('index.html') # Return here your Html Page (with the form)


@app.route('/', methods=['GET','POST']) # When the user click to the "submit" button
def my_form_post():
    
    document = docx.Document()

# First we put into variables all the data that the user entered
    fname= request.form.get('firstname')
    lname = request.form.get('lastname')
    country = request.form.get('country')
    
    
    document.add_heading("My Name is {}".format(fname), 0)

    docname = 'Demo{}.docx'.format(fname)

    document.add_paragraph("I have been to {}.".format(country))


    document.save(docname)
    
    
    #return render_template('index.html') # Return here your Html Page (with the form)
    #return 'Hello, World!'
    return send_file(docname, as_attachment=True)
    #return send_file('root.docx',name,as_attachment=True, attachment_filename=os.path.basename(name))


if __name__ == "__main__":
    app.run(host='0.0.0.0',port='5001')
